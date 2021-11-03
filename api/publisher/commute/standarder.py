from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_TAB_ALIGNMENT
from docx.shared import Cm, Pt, RGBColor
from re import sub
from datetime import date, timedelta

from publisher.commute import fstring, freplace


# Substitute text fragment
def find_replace(document):
    for p in document.paragraphs:
        for r in p.runs:
            r.text = sub(r"[\r\n]", " ", r.text)
            r.text = sub(r"  +", " ", r.text)
            r.text = sub(r" ,+", ",", r.text)
            r.text = sub(r"(–|- -)", "-", r.text)
            r.text = sub(r"(\.?-\.?)", "-", r.text)
            r.text = sub(r"(°|º)", "º", r.text)
            r.text = sub(r"(\d)[oO]", "\g<1>º", r.text)
            r.text = sub(r"(\d)[aA]", "\g<1>ª", r.text)
            r.text = sub(r"(a|A)rt\.", "\g<1>rtigo", r.text)
            r.text = sub(r"(n\.º)", "nº", r.text)
            r.text = sub(r"(nº) ", "\g<1>\u00a0", r.text)
            r.text = sub(r"(-) ", "\g<1>\u00a0", r.text)


# Extract the text of initial document
def extract_text(document, document_default, bold, italic, underline, format_allowed, condensation):
    # Document flowing text

    texts_flowing = ''

    if format_allowed == '2':
        index = 0
        texts_flowing = document_default.paragraphs[0]
    elif format_allowed in ['1', '2']:
        index = 1
        texts_flowing = document_default.add_paragraph()
    else:
        index = 3
        texts_flowing = document_default.add_paragraph()

    texts_flowing.paragraph_format.tab_stops.add_tab_stop(document_default.sections[0].page_width - Cm(0.05), WD_TAB_ALIGNMENT.RIGHT)

    for idx, val in enumerate(document.paragraphs):
        if index > 0:
            if val.text != '':
                index = index - 1
        else:
            if val.text != '':
                run_text(val, texts_flowing, bold, italic, underline, condensation)

def insert_date(document_default, format_parameters):
    # Insert publication date(s)

    if format_parameters['name_section'] == "Terceiros - DC - MT":
        date_today = date.today()
        document_default.paragraphs[-1].add_run('\t')
        print(format_parameters['days'])
        if format_parameters['days'] == "1":
            document_default.paragraphs[-1].add_run((date_today + timedelta(days=1)).strftime('%d/%m/%y'))
        elif format_parameters['days'] == "2":
            document_default.paragraphs[-1].add_run((date_today + timedelta(days=1)).strftime('%d/%m') + " e " +
                                                    (date_today + timedelta(days=2)).strftime('%d/%m/%y'))
        elif format_parameters['days'] == "3":
            document_default.paragraphs[-1].add_run((date_today + timedelta(days=1)).strftime('%d/%m') + ", " +
                                                    (date_today + timedelta(days=2)).strftime('%d/%m') + " e " +
                                                    (date_today + timedelta(days=3)).strftime('%d/%m/%y'))

# Read and apply bold, italic and underlines
def run_text(paragraph, texts_flowing, bold, italic, underline, condensation):
    if bold or italic or underline:
        if len(texts_flowing.runs) == 0:
            texts_flowing.add_run(style=condensation)
            texts_flowing.runs[-1].font.color.rgb = RGBColor(0x00, 0x00, 0x00)
        if len(paragraph.runs) > 0:
            for r in paragraph.runs:
                r.text = sub('  +', ' ', r.text)
                if r.text != '' and r.text != ' ':
                    if r.bold == texts_flowing.runs[-1].bold and r.italic == texts_flowing.runs[-1].italic and r.underline == texts_flowing.runs[-1].underline:
                        #if texts_flowing.runs[-1].text[-1] == ' ' and r.text[0] == ' ':
                        if len(texts_flowing.runs[-1].text) > 0 and texts_flowing.runs[-1].text[-1] == ' ' and r.text[0] == ' ':
                            texts_flowing.runs[-1].add_text(r.text[1:])
                        else:
                            texts_flowing.runs[-1].add_text(r.text)
                    else:
                        texts_block = texts_flowing.add_run(r.text, style=condensation)
                        texts_block.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
                        if bold:
                            texts_block.bold = r.bold
                        if italic:
                            texts_block.italic = r.italic
                        if underline:
                            texts_block.underline = r.underline
                else:
                    if len(texts_flowing.runs) > 1:
                        texts_flowing.runs[-1].add_text(' ')

    else:
        texts_flowing.add_run(paragraph.text, style=condensation)

# Define document type formatting | Head lines
def read_text(document, document_default, format_allowed):
    if format_allowed == "0":
        index = 3
        for idx, val in enumerate(document.paragraphs):
            if index > 0:
                val.text = sub('  +', ' ', val.text)
                if val.text != '' and val.text != ' ':
                    index = index - 1
                    if document_default.paragraphs[0].text == '':
                        document_default.paragraphs[0].text = val.text
                    else:
                        document_default.add_paragraph(val.text)
            else:
                break
    elif format_allowed == "1":
        index = 1
        for idx, val in enumerate(document.paragraphs):
            if index > 0:
                val.text = sub('  +', ' ', val.text)
                if val.text != '' and val.text != ' ':
                    index = index - 1
                    document_default.paragraphs[0].text = val.text
            else:
                break

# Format document head
def doc_format_head(document_default, format_allowed, font_size_company,
                    font_leading_company, bold, texts, i, format_parameters):
    if format_allowed == "0":

        # Center first 3 lines
        if i < 3:
            document_default.paragraphs[i].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            texts.bold = bold
            #document_default.paragraphs[i].paragraph_format.text. = Pt(1.1)
            if format_parameters['name_section'] == "Particulares - DOBA - BA":
                document_default.paragraphs[i].paragraph_format.space_after = Pt(1.1)

        # Manipulating first line
        if i == 0:
            texts.font.size = Pt(font_size_company)
            document_default.paragraphs[i].paragraph_format.line_spacing = Pt(font_leading_company)
            #if format_parameters['name_section'] == "Particulares - DOBA - BA":

        # Manipulating third line
        if i == 2 and format_parameters['name_section'] == "Particulares - DOBA - BA":
            document_default.paragraphs[i].paragraph_format.line_spacing = Pt(13)
            texts.font.size = Pt(12)

        # Manipulating second line
        if i == 1:
            texts.bold = False
            #if format_parameters['name_section'] == "Particulares - DOBA - BA":
                #texts.line_spacing = Pt(10)
                # texts.line_spacing =

    elif format_allowed == "1" and i < 1:
        document_default.paragraphs[i].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        texts.bold = bold

# Format document content
def doc_format(document_default, font_name, font_size, font_size_company, font_leading, font_leading_company, 
               bold, format_allowed, edge, format_parameters):
    for i in range(len(document_default.paragraphs)):
        document_default.paragraphs[i].paragraph_format.space_after = Pt(0)
        document_default.paragraphs[i].paragraph_format.space_before = Pt(0)
        document_default.paragraphs[i].alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        # Borders
        if edge is True:
            border_size = 4
            border_space = 1
            border_indent = 0.05
            if format_parameters['name_section'] == "Terceiros - DOBA - BA":
                border_space = 4
                border_indent = 1
            document_default.paragraphs[i].paragraph_format.left_indent = Cm(border_indent)
            document_default.paragraphs[i].paragraph_format.right_indent = Cm(border_indent)
            document_default.paragraphs[i].paragraph_format.border_top(border_size, border_space)
            document_default.paragraphs[i].paragraph_format.border_bottom(border_size, border_space)
            document_default.paragraphs[i].paragraph_format.border_left(border_size, border_space)
            document_default.paragraphs[i].paragraph_format.border_right(border_size, border_space)

        # Special format
        for texts in document_default.paragraphs[i].runs:
            texts.font.name = font_name
            texts.font.size = Pt(font_size)
            document_default.paragraphs[i].paragraph_format.line_spacing = Pt(font_leading)
            doc_format_head(document_default, format_allowed, font_size_company,
                            font_leading_company, bold, texts, i, format_parameters)

# Receive initial document, format and return a standardized document
def doc_default(format_parameters):

    print(fstring.message["info"]['standard_start'])

    document = Document(format_parameters['document_name']+'.docx')
    document_default = Document(fstring.paths['model'])

    # Document size
    document_default.sections[0].page_width = Cm(format_parameters['column'])
    document_default.sections[0].page_height = Cm(format_parameters['height'])

    # Document margins
    document_default.sections[0].top_margin = Cm(0)
    document_default.sections[0].left_margin = Cm(0)
    document_default.sections[0].right_margin = Cm(0)
    document_default.sections[0].bottom_margin = Cm(0)

    # Document text read
    read_text(document, document_default, format_parameters['format_allowed'])

    # Document extract
    extract_text(document, document_default, format_parameters['bold'], format_parameters['italic'], 
                    format_parameters['underline'], format_parameters['format_allowed'], format_parameters['condensation'])

    # F/R
    find_replace(document_default)

    # Insert tab for date
    insert_date(document_default, format_parameters)

    # Document formatting
    doc_format(document_default, format_parameters['font_name'], format_parameters['font_size'], format_parameters['font_size_company'],
                format_parameters['font_leading'], format_parameters['font_leading_company'], format_parameters['bold'], 
                format_parameters['format_allowed'], format_parameters['edge'], format_parameters)

    # Document save
    document_default.save(format_parameters['document_name']+'.docx')
    chars_count = 0
    for p in document_default.paragraphs:
        chars_count = chars_count + len(p.text)
    print(fstring.message["info"]['docx_success'])
    return chars_count
