import logging
from datetime import date, timedelta
from re import sub

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_TAB_ALIGNMENT
from docx.shared import Cm, Pt, RGBColor

from publisher.commute import fstring

logging.basicConfig(level=logging.INFO)


def find_replace(format_parameter):
    for p in format_parameter['document'].paragraphs:
        for r in p.runs:
            r.text = sub(r"[\r\n]", " ", r.text)
            r.text = sub(r"  +", " ", r.text)
            r.text = sub(r" ,+", ",", r.text)
            r.text = sub(r"(–|- -)", "-", r.text)
            r.text = sub(r"(\.?-\.?)", "-", r.text)
            r.text = sub(r"(°|º)", "º", r.text)
            r.text = sub(r"(\d)[oO]", "\\g<1>º", r.text)
            r.text = sub(r"(\d)[aA]", "\\g<1>ª", r.text)
            r.text = sub(r"(a|A)rt\.", "\\g<1>rtigo", r.text)
            r.text = sub(r"(n\.º)", "nº", r.text)
            r.text = sub(r"(nº) ", "\\g<1>\u00a0", r.text)
            r.text = sub(r"(-) ", "\u00a0\\g<1> ", r.text)
            r.text = sub(r" \u00a0+", "\u00a0", r.text)
    for p in format_parameter['docx_default'].paragraphs:
        for r in p.runs:
            r.text = sub(r"[\r\n]", " ", r.text)
            r.text = sub(r"  +", " ", r.text)
            r.text = sub(r" ,+", ",", r.text)
            r.text = sub(r"(–|- -)", "-", r.text)
            r.text = sub(r"(\.?-\.?)", "-", r.text)
            r.text = sub(r"(°|º)", "º", r.text)
            r.text = sub(r"(\d)[oO]", "\\g<1>º", r.text)
            r.text = sub(r"(\d)[aA]", "\\g<1>ª", r.text)
            r.text = sub(r"(a|A)rt\.", "\\g<1>rtigo", r.text)
            r.text = sub(r"(n\.º)", "nº", r.text)
            r.text = sub(r"(nº) ", "\\g<1>\u00a0", r.text)
            r.text = sub(r" ?(-) ", "\u00a0\\g<1> ", r.text)
            r.text = sub(r" \u00a0+", "\u00a0", r.text)


def extract_text(format_parameters):
    format_parameters['texts_flowing'] = ''

    if format_parameters['format_allowed'] == '2':
        index: int = 0
        format_parameters['texts_flowing'] = format_parameters['docx_default'].paragraphs[0]
    elif format_parameters['format_allowed'] in ['1', '2']:
        index: int = 1
        format_parameters['texts_flowing'] = format_parameters['docx_default'].add_paragraph()
    else:
        index: int = 3
        format_parameters['texts_flowing'] = format_parameters['docx_default'].add_paragraph()

    format_parameters['texts_flowing'].paragraph_format.tab_stops.add_tab_stop(
        format_parameters['docx_default'].sections[0].page_width - Cm(0.05),
        WD_TAB_ALIGNMENT.RIGHT)

    for idx, format_parameters['paragraph'] in enumerate(format_parameters['document'].paragraphs):
        if index > 0:
            if format_parameters['paragraph'].text != '':
                index: int = index - 1
        else:
            if format_parameters['paragraph'].text != '':
                run_text(format_parameters)


def insert_date(format_parameters):
    docx_default = format_parameters['docx_default']
    if format_parameters['name_section'] in fstring.newspapers['insert_date']:
        date_today: date = date.today()
        docx_default.paragraphs[-1].add_run('\t')
        logging.info(format_parameters['days'])

        if format_parameters['days'] == "1":
            docx_default.paragraphs[-1].add_run(
                (date_today + timedelta(days=1)).strftime('%d/%m/%y')
            )
        elif format_parameters['days'] == "2":
            docx_default.paragraphs[-1].add_run(
                (date_today + timedelta(days=1)).strftime('%d/%m') + " e " +
                (date_today + timedelta(days=2)).strftime('%d/%m/%y')
            )
        elif format_parameters['days'] == "3":
            docx_default.paragraphs[-1].add_run(
                (date_today + timedelta(days=1)).strftime('%d/%m') + ", " +
                (date_today + timedelta(days=2)).strftime('%d/%m') + " e " +
                (date_today + timedelta(days=3)).strftime('%d/%m/%y')
            )


def set_color(format_parameters):
    format_parameters['texts_flowing'].runs[-1].font.color.rgb = RGBColor(0x00, 0x00, 0x00)


def run_text(format_parameters):
    if format_parameters['bold'] or format_parameters['italic'] or format_parameters['underline']:
        if len(format_parameters['texts_flowing'].runs) == 0:
            format_parameters['texts_flowing'].add_run(style=format_parameters['condensation'])
            set_color(format_parameters)
        if len(format_parameters['paragraph'].runs) > 0:
            for r in format_parameters['paragraph'].runs:
                r.text = sub('  +', ' ', r.text)
                if r.text != '' and r.text != ' ':
                    if r.bold == format_parameters['texts_flowing'].runs[-1].bold and r.italic == \
                            format_parameters['texts_flowing'].runs[
                                -1].italic and r.underline == format_parameters['texts_flowing'].runs[-1].underline:
                        if len(format_parameters['texts_flowing'].runs[-1].text) > 0 and \
                                format_parameters['texts_flowing'].runs[-1].text[-1] == ' ' and r.text[0] == ' ':
                            format_parameters['texts_flowing'].runs[-1].add_text(r.text[1:])
                        else:
                            format_parameters['texts_flowing'].runs[-1].add_text(r.text)
                    else:
                        texts_block = format_parameters['texts_flowing'].add_run(r.text, style=format_parameters[
                            'condensation'])
                        set_color(format_parameters)
                        if format_parameters['bold']:
                            texts_block.bold = r.bold
                        if format_parameters['italic']:
                            texts_block.italic = r.italic
                        if format_parameters['underline']:
                            texts_block.underline = r.underline
                else:
                    if len(format_parameters['texts_flowing'].runs) > 1:
                        format_parameters['texts_flowing'].runs[-1].add_text(' ')

    else:
        format_parameters['texts_flowing'].add_run(format_parameters['paragraph'].text,
                                                   style=format_parameters['condensation'])


def read_text(format_parameters):
    if format_parameters['format_allowed'] == "0":
        format_parameters['text_index'] = 3
        for idx, format_parameters['initial_paragraph'] in enumerate(format_parameters['document'].paragraphs):
            if format_parameters['text_index'] > 0:
                format_parameters['initial_paragraph'].text = sub('  +', ' ',
                                                                  format_parameters['initial_paragraph'].text)
                if format_parameters['initial_paragraph'].text != '' and format_parameters[
                    'initial_paragraph'].text != ' ':
                    format_parameters['text_index'] = format_parameters['text_index'] - 1
                    if format_parameters['docx_default'].paragraphs[0].text == '':
                        format_parameters['docx_default'].paragraphs[0].text = format_parameters[
                            'initial_paragraph'].text
                    else:
                        format_parameters['docx_default'].add_paragraph(format_parameters['initial_paragraph'].text)
            else:
                break
    elif format_parameters['format_allowed'] == "1":
        format_parameters['text_index'] = 1
        for idx, format_parameters['initial_paragraph'] in enumerate(format_parameters['document'].paragraphs):
            if format_parameters['text_index'] > 0:
                format_parameters['initial_paragraph'].text = sub('  +', ' ',
                                                                  format_parameters['initial_paragraph'].text)
                if format_parameters['initial_paragraph'].text != '' and format_parameters['initial_paragraph'].text != ' ':
                    format_parameters['text_index'] = format_parameters['text_index'] - 1
                    format_parameters['docx_default'].paragraphs[0].text = format_parameters['initial_paragraph'].text
            else:
                break


class Format:

    @staticmethod
    def format_text_head(i, format_parameters):
        if format_parameters['format_allowed'] == "0":

            if i < 3:
                format_parameters['docx_default'].paragraphs[i].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                format_parameters['texts'].bold = format_parameters['bold']
                if format_parameters['name_section'] in fstring.newspapers['special_format_head_spaced']:
                    format_parameters['docx_default'].paragraphs[i].paragraph_format.space_after = Pt(1.1)

            if i == 0:
                format_parameters['texts'].font.size = Pt(format_parameters['font_size_company'])
                format_parameters['docx_default'].paragraphs[i].paragraph_format.line_spacing = Pt(
                    format_parameters['font_leading_company'])

            if i == 2 and format_parameters['name_section'] in fstring.newspapers['special_format_head_spaced']:
                format_parameters['docx_default'].paragraphs[i].paragraph_format.line_spacing = Pt(13)
                format_parameters['texts'].font.size = Pt(12)

            if i == 1:
                format_parameters['texts'].bold = False

        elif format_parameters['format_allowed'] == "1" and i < 1:
            format_parameters['docx_default'].paragraphs[i].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            format_parameters['texts'].bold = format_parameters['bold']

    @staticmethod
    def format_document_size(format_parameters: dict):
        column = format_parameters['column']
        format_parameters['docx_default'].sections[0].page_width = Cm(column)
        format_parameters['docx_default'].sections[0].page_height = Cm(format_parameters['height'])

    @staticmethod
    def format_margins(format_parameters: dict):
        document_margins = [0, 0, 0, 0]
        if format_parameters['name_section'] in fstring.newspapers['special_format_a4'] and \
                format_parameters['number_column'] == 1:
            document_margins[2] = 13
        elif format_parameters['name_section'] in fstring.newspapers['special_format_a4'] \
                and format_parameters['number_column'] == 2:
            document_margins[2] = 4
        format_parameters['docx_default'].sections[0].top_margin = Cm(document_margins[0])
        format_parameters['docx_default'].sections[0].left_margin = Cm(document_margins[1])
        format_parameters['docx_default'].sections[0].right_margin = Cm(document_margins[2])
        format_parameters['docx_default'].sections[0].bottom_margin = Cm(document_margins[3])

    @staticmethod
    def format_page(format_parameters: dict):
        for i in range(len(format_parameters['docx_default'].paragraphs)):
            format_parameters['docx_default'].paragraphs[i].paragraph_format.space_after = Pt(0)
            format_parameters['docx_default'].paragraphs[i].paragraph_format.space_before = Pt(0)
            format_parameters['docx_default'].paragraphs[i].alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

            add_edge(format_parameters, i)

            for format_parameters['texts'] in format_parameters['docx_default'].paragraphs[i].runs:
                format_parameters['texts'].font.name = format_parameters['font_name']
                format_parameters['texts'].font.size = Pt(format_parameters['font_size'])
                format_parameters['docx_default'].paragraphs[i].paragraph_format.line_spacing = Pt(
                    format_parameters['font_leading'])
                Format.format_text_head(i, format_parameters)


def add_edge(format_parameters: dict, i):
    if format_parameters['edge'] is True:
        border_size: int = 4
        border_space: int = 1
        border_indent: float = 0.05
        if format_parameters['name_section'] == "Terceiros - DOBA - BA":
            border_space: int = 4
            border_indent: int = 1
        format_parameters['docx_default'].paragraphs[i].paragraph_format.left_indent = Cm(border_indent)
        format_parameters['docx_default'].paragraphs[i].paragraph_format.right_indent = Cm(border_indent)
        format_parameters['docx_default'].paragraphs[i].paragraph_format.border_top(border_size, border_space)
        format_parameters['docx_default'].paragraphs[i].paragraph_format.border_bottom(border_size,
                                                                                       border_space)
        format_parameters['docx_default'].paragraphs[i].paragraph_format.border_left(border_size, border_space)
        format_parameters['docx_default'].paragraphs[i].paragraph_format.border_right(border_size, border_space)


def special_adjust(format_parameters: dict):
    if format_parameters['name_section'] in fstring.newspapers['special_format_a4']:
        docx_default = format_parameters['docx_default']
        column: int = 21
        docx_default.sections[0].page_width = Cm(column)
        document_right_margin: int = 0
        if format_parameters['number_column'] == "1":
            document_right_margin: int = 13
        elif format_parameters['number_column'] == "2":
            document_right_margin: int = 4
        docx_default.sections[0].right_margin = Cm(document_right_margin)


def save_document(format_parameters: dict):
    format_parameters['docx_default'].save(format_parameters['document_name'] + '.docx')
    chars_count: int = 0
    for p in format_parameters['docx_default'].paragraphs:
        format_parameters['chars_count'] = chars_count + len(p.text)
    logging.info(fstring.message["info"]['docx_success'])


def standardizer_docx(format_parameters: dict):
    logging.info(fstring.message["info"]['standard_start'])
    format_parameters['document'] = Document(format_parameters['document_name'] + '.docx')
    format_parameters['docx_default'] = Document(fstring.paths['model'])

    Format.format_document_size(format_parameters)
    Format.format_margins(format_parameters)
    read_text(format_parameters)
    extract_text(format_parameters)
    find_replace(format_parameters)
    insert_date(format_parameters)
    Format.format_page(format_parameters)
    special_adjust(format_parameters)
    save_document(format_parameters)

    return format_parameters['chars_count']
